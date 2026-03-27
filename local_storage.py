import os
from docx import Document
from docx.shared import Pt, RGBColor
from config import SAVE_DIRECTORY_PATH, logger

def create_and_save_report(mail_data, ai_result):
    """
    사용자가 텔레그램 방에서 '/save' 라는 특별한 도장(명령어)을 쾅! 찍어주셨을 때만 
    단 한 번만 작동하는 충직한 비서 같은 함수입니다.
    
    분석이 끝난 메일 내용과 AI의 훌륭한 요약을 모아서, 가상 공간이 아닌 
    '내 진짜 컴퓨터' 안에 깔끔한 '워드 문서(.docx)' 파일로 고이 접어 보관해줍니다.
    절대 자기 마음대로 문서를 찍어내서 컴퓨터 용량을 잡아먹지 않도록 통제됩니다.
    """
    
    # 1. 파일들을 예쁘게 모아둘 서랍장(폴더)이 내 컴퓨터에 진짜 있는지 확인하고, 없으면 그때 새로 하나 짜맞춥니다.
    if not os.path.exists(SAVE_DIRECTORY_PATH):
        try:
            os.makedirs(SAVE_DIRECTORY_PATH)
            logger.info(f"문서를 보관할 튼튼한 금고(새 폴더)를 만들었습니다: {SAVE_DIRECTORY_PATH}")
        except Exception as e:
            logger.error(f"저장 폴더를 만드는 데 문제가 생겼습니다. 튼튼하지 않은 거 같습니다: {e}")
            return False, "폴더 생성 실패"

    # 2. 문구점에서 공책을 살 때 겉표지에 반이랑 이름을 예쁘게 적어두듯, 
    # 날짜와 종류를 조합해서 겹치지 않는 파일 이름을 정해줍니다. (예: 20260326_영업_A사견적서.docx)
    
    # 윈도우(Windows) 운영체제는 제목에 \ / : * ? " < > | 같은 특수문자를 아주 싫어하므로 안전하게 걸러냅니다. (보안 및 에러 방지)
    raw_subject = mail_data.get('subject', '제목없음')
    safe_title = "".join(c for c in raw_subject if c.isalnum() or c in (' ', '_', '-')).strip()
    
    # "2026-03-26 14:22:30" 같은 복잡한 시간에서 년, 월, 일만 깔끔하게 오려냅니다.
    short_date = mail_data.get('date', '날짜미상')[:10].replace(" ", "").replace("-", "") 
    category = ai_result.get('category', '미분류')
    
    # 최종적으로 정해진 파일의 이름과, 내 컴퓨터 안의 진짜 경로입니다!
    file_name = f"{short_date}_{category}_{safe_title[:20]}.docx"
    file_path = os.path.join(SAVE_DIRECTORY_PATH, file_name)

    try:
        # 3. 이제 새하얀 빈 백지 워드 문서를 한 장 펼칩니다.
        doc = Document()
        
        # [문서의 제일 윗면 큰 제목]
        doc.add_heading(f"📧 이메일 인공지능 분석 보고서: {raw_subject}", level=1)
        
        # [기본 정보 표] 누가 언제 보냈는지 헷갈리지 않게 깔끔한 2칸짜리 표 안에 적어둡니다.
        doc.add_heading("1. 메일 기본 정보", level=2)
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        
        row0 = table.rows[0].cells
        row0[0].text = "보냈던 사람"
        row0[1].text = mail_data.get('sender', '알 수 없음')
        
        row1 = table.rows[1].cells
        row1[0].text = "받았던 일시"
        row1[1].text = mail_data.get('date', '알 수 없음')
        
        # [AI 요약 및 실무 조언]
        doc.add_heading("2. 전체 흐름 요약 및 전문가 조언", level=2)
        p_summary = doc.add_paragraph()
        p_summary.add_run("[ 주고받은 과거 흐름 요약 ]\n").bold = True
        p_summary.add_run(ai_result.get('summary', '요약 내용이 없습니다.') + "\n\n")
        
        p_advice = doc.add_paragraph()
        p_advice.add_run("[ 실무진 조언 ]\n").bold = True
        p_advice.add_run(ai_result.get('advice', '조언 내용이 없습니다.'))
        
        # [원본 이메일] 만약 AI 요약을 봤는데도 헷갈리면 제일 밑에서 실제로 온 편지를 그대로 볼 수 있게 해줍니다.
        doc.add_heading("3. 원문 텍스트 편지 원본", level=2)
        p_original_title = doc.add_paragraph()
        p_original_title.add_run("(아래는 이메일 전체 원본 글씨만 뽑아온 내용입니다.)").italic = True
        
        p_body = doc.add_paragraph(mail_data.get('body', '본문 내용 없음'))
        # 원본 글씨는 중요도가 떨어지니 크기를 살짝 줄이고 회색 빛깔로 바꿔 눈에 조금 부담을 덜어줍니다.
        for run in p_body.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # 4. 방금 정성껏 옮겨 적은 문서를 컴퓨터 서랍장에 넣고 '딸깍' 잠가 안전하게 저장합니다!
        doc.save(file_path)
        logger.info(f"워드 보고서가 너무 예쁘게 잘 완성되었습니다: {file_path}")
        
        return True, file_path

    except Exception as e:
        logger.error(f"보고서를 예쁘게 정리하다 예기치 않은 볼펜 잉크 번짐(에러)이 쓰였어요: {e}")
        return False, f"문서 생성 에러: {e}"
